import pandas as pd
from docx import Document
import os

# Folder path where your Excel and CSV files are stored
folder_path = r"C:/Users/Shreyash Bhardwaj/Desktop/Excel_Extraction"

# List of files to process
excel_files = [
    "Muzaffarpur.xls",
    "Patna.xls",
    # "Bardhaman.xlt.xls",
    # "Durgapur.xlsx",
    # "Howrah.xlt.xls",
    # "Kharagpur.xls",
    "Ranchi.xls"
]

# Process each file
for file in excel_files:
    try:
        ext = os.path.splitext(file)[1].lower()
        file_path = os.path.join(folder_path, file)  # Construct full file path first

        # Read based on extension
        if ext == '.xlsx':
            df = pd.read_excel(file_path, engine='openpyxl')
        elif ext == '.xls':
            df = pd.read_excel(file_path, engine='xlrd')
        elif ext == '.csv':
            df = pd.read_csv(file_path)
        else:
            print(f"⚠️ Skipping unsupported file type: {file}")
            continue

        # Extract Column F (index 5) from excel
        column_data = df.iloc[:, 5].dropna().astype(str).tolist()

        # Use the below one for when using csv
        # column_data = df['Email(required)'].dropna().astype(str).tolist()
        comma_separated = ', '.join(column_data)

        # Create and save the Word document
        doc = Document()
        doc.add_heading(f'Data from Column F - {file}', level=1)
        doc.add_paragraph(comma_separated)

        output_name = os.path.splitext(file)[0] + ".docx"
        output_path = os.path.join(folder_path, output_name)

        doc.save(output_path)
        print(f"✅ Saved: {output_name}")

    except Exception as e:
        print(f"⚠️ Failed to process {file}: {e}")
