import pandas as pd
from docx import Document
import os
import re

# === Config ===
excel_file_path = "ARN MAster file 7 Cities.xlsx"  # Your file
output_folder = "city_email_docs"

# === Create output folder if not exists ===
os.makedirs(output_folder, exist_ok=True)

# === Load the Excel file ===
df = pd.read_excel(excel_file_path)

# === Rename columns for clarity ===
df.columns = df.columns.str.strip()
df = df.rename(columns={
    df.columns[1]: "ARN",    # Column B
    df.columns[4]: "Email",  # Column E
    df.columns[5]: "City"    # Column F
})

# === Filter data ===
filtered_df = df[(df["ARN"] >= 150000) & df["Email"].notna() & df["City"].notna()]

# === Normalize city names to lower-case for fuzzy matching ===
filtered_df["City_clean"] = filtered_df["City"].str.lower().str.strip()

# === Extract all unique base city names from the data ===
unique_cities = set()
for city in filtered_df["City_clean"]:
    # Get just the main word like "ahmedabad" from "ahmedabad city"
    match = re.search(r"(chennai|ahmedabad|patna|ranchi|muzaffarpur|durgapur|howrah|kharagpur|bardhaman|bangalore|mumbai|pune|delhi|kolkata)", city)
    if match:
        unique_cities.add(match.group(1))

# === For each base city, collect all matching rows ===
for base_city in unique_cities:
    matching_rows = filtered_df[filtered_df["City_clean"].str.contains(base_city)]

    emails = matching_rows["Email"].dropna().astype(str).unique()

    if len(emails) == 0:
        continue

    doc = Document()
    doc.add_heading(f"Email List - {base_city.title()}", level=1)
    doc.add_paragraph(', '.join(emails))

    output_path = os.path.join(output_folder, f"{base_city}.docx")
    doc.save(output_path)
    print(f"✅ Saved: {output_path}")
